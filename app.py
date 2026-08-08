import json
import uuid
import os
import io
import re
from datetime import datetime

from fastapi import FastAPI, Request
from fastapi.responses import HTMLResponse, StreamingResponse, JSONResponse, Response
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

from pipeline import run_research_pipeline_stream
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor

# ── App setup ────────────────────────────────────────────────

app = FastAPI(title="SYNAPSE", description="AI Research Assistant")

# CORS — allow development server and deployed frontend domains
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

HISTORY_FILE = os.path.join(os.path.dirname(__file__), "research_history.json")
FRONTEND_DIST = os.path.join(os.path.dirname(__file__), "frontend", "dist")

if os.path.exists(FRONTEND_DIST):
    app.mount("/assets", StaticFiles(directory=os.path.join(FRONTEND_DIST, "assets")), name="assets")

    @app.get("/")
    async def serve_spa():
        return HTMLResponse(open(os.path.join(FRONTEND_DIST, "index.html"), encoding="utf-8").read())


# ── Pydantic models ─────────────────────────────────────────

class PDFRequest(BaseModel):
    report: str
    topic: str = "Research Report"


class DocxRequest(BaseModel):
    report: str
    topic: str = "Research Report"


# ── History helpers ──────────────────────────────────────────

def load_history():
    """Load research history from JSON file."""
    if not os.path.exists(HISTORY_FILE):
        return []
    try:
        with open(HISTORY_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except (json.JSONDecodeError, IOError):
        return []


def save_history(history):
    """Save research history to JSON file."""
    with open(HISTORY_FILE, "w", encoding="utf-8") as f:
        json.dump(history, f, ensure_ascii=False, indent=2)


def add_to_history(topic, state):
    """Append a completed research to history."""
    history = load_history()
    entry = {
        "id": str(uuid.uuid4()),
        "topic": topic,
        "timestamp": datetime.now().isoformat(),
        "report": state.get("report", ""),
        "feedback": state.get("feedback", ""),
        "search_results": state.get("search_results", ""),
        "scraped_content": state.get("scraped_content", ""),
    }
    history.insert(0, entry)  # newest first
    save_history(history)
    return entry["id"]


# ── PDF generation ───────────────────────────────────────────

class ResearchPDF(FPDF):
    """Custom PDF with header/footer for research reports."""

    def header(self):
        self.set_font("Helvetica", "B", 10)
        self.set_text_color(100, 100, 100)
        self.cell(0, 8, "SYNAPSE Research Report", align="L")
        self.ln(10)

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "", 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")


def _sanitize(text):
    """Remove characters unsupported by the built-in Helvetica font."""
    replacements = {
        '\u2013': '-', '\u2014': '--', '\u2018': "'", '\u2019': "'",
        '\u201c': '"', '\u201d': '"', '\u2026': '...', '\u2022': '-',
        '\u2010': '-', '\u2011': '-', '\u2012': '-', '\u00a0': ' ',
        '\u200b': '', '\u200c': '', '\u200d': '', '\ufeff': '',
        '\u2713': 'v', '\u2717': 'x', '\u2192': '->', '\u2190': '<-',
        '\u00b7': '-', '\u25cf': '-', '\u25cb': 'o', '\u2605': '*',
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text.encode('latin-1', errors='replace').decode('latin-1')


def markdown_to_pdf(report_text, topic="Research Report"):
    """Convert markdown report text to a styled PDF."""
    pdf = ResearchPDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(30, 30, 30)
    pdf.multi_cell(0, 10, _sanitize(topic))
    pdf.ln(4)

    # Timestamp
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(130, 130, 130)
    pdf.cell(0, 6, f"Generated on {datetime.now().strftime('%B %d, %Y at %H:%M')}")
    pdf.ln(10)

    # Line separator
    pdf.set_draw_color(200, 200, 200)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(6)

    # Parse and render markdown lines
    for line in report_text.split("\n"):
        stripped = line.strip()

        if not stripped:
            pdf.ln(3)
            continue

        if stripped.startswith("### "):
            pdf.ln(4)
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_text_color(50, 50, 50)
            pdf.multi_cell(0, 7, _sanitize(stripped[4:]))
            pdf.ln(2)

        elif stripped.startswith("## "):
            pdf.ln(5)
            pdf.set_font("Helvetica", "B", 14)
            pdf.set_text_color(40, 40, 40)
            pdf.multi_cell(0, 8, _sanitize(stripped[3:]))
            pdf.ln(2)

        elif stripped.startswith("# "):
            pdf.ln(6)
            pdf.set_font("Helvetica", "B", 16)
            pdf.set_text_color(30, 30, 30)
            pdf.multi_cell(0, 9, _sanitize(stripped[2:]))
            pdf.ln(3)

        elif stripped.startswith("- ") or stripped.startswith("* "):
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(60, 60, 60)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', stripped[2:])
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            pdf.cell(6)
            pdf.multi_cell(0, 6, "  - " + _sanitize(text))
            pdf.ln(1)

        elif re.match(r'^\d+\.\s', stripped):
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(60, 60, 60)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', stripped)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            pdf.cell(6)
            pdf.multi_cell(0, 6, _sanitize(text))
            pdf.ln(1)

        else:
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(60, 60, 60)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', stripped)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            pdf.multi_cell(0, 6, _sanitize(text))
            pdf.ln(1)

    return pdf.output()


# ── Routes ───────────────────────────────────────────────────

# ── Serve React build in production ──────────────────────────

FRONTEND_DIST = os.path.join(os.path.dirname(__file__), "frontend", "dist")


@app.get("/", response_class=HTMLResponse)
async def index():
    index_path = os.path.join(FRONTEND_DIST, "index.html")
    if os.path.exists(index_path):
        with open(index_path, "r", encoding="utf-8") as f:
            return HTMLResponse(content=f.read())
    return HTMLResponse(
        content="<h1>SYNAPSE API</h1><p>Frontend not built. Run <code>npm run build</code> in frontend/.</p>"
    )


@app.get("/run")
async def run_pipeline(topic: str = ""):
    """SSE endpoint — streams pipeline step events as JSON."""
    topic = topic.strip()

    if not topic:
        return JSONResponse({"error": "No topic provided"})

    def generate():
        try:
            for event in run_research_pipeline_stream(topic):
                if event.get("step") == "complete" and "state" in event:
                    entry_id = add_to_history(topic, event["state"])
                    event["history_id"] = entry_id

                yield f"data: {json.dumps(event)}\n\n"
        except Exception as e:
            yield f'data: {json.dumps({"error": str(e)})}\n\n'

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        }
    )


def markdown_to_docx(report_text, topic="Research Report"):
    """Convert markdown report text to a beautifully formatted Word (.docx) document."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Document Header Title
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(topic)
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(30, 41, 59)

    # Subtitle Metadata
    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run(f"SYNAPSE Autonomous AI Research Report  |  Generated on {datetime.now().strftime('%B %d, %Y at %H:%M')}")
    sub_run.font.name = 'Calibri'
    sub_run.font.size = Pt(9.5)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    for line in report_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(stripped[2:])
            run.font.name = 'Calibri'
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(124, 108, 240)

        elif stripped.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped[3:])
            run.font.name = 'Calibri'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(30, 41, 59)

        elif stripped.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(stripped[4:])
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(71, 85, 105)

        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', stripped[2:])
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            run = p.add_run(text)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(51, 65, 85)

        elif re.match(r'^\d+\.\s', stripped):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(3)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', stripped)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            clean_text = re.sub(r'^\d+\.\s*', '', text)
            run = p.add_run(clean_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(51, 65, 85)

        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', stripped)

            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    run = p.add_run(part)
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(51, 65, 85)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


@app.post("/download-pdf")
async def download_pdf(data: PDFRequest):
    """Generate and return a PDF from the report text."""
    if not data.report:
        return JSONResponse({"error": "No report provided"}, status_code=400)

    pdf_bytes = markdown_to_pdf(data.report, data.topic)
    buffer = io.BytesIO(pdf_bytes)
    buffer.seek(0)

    safe_name = re.sub(r'[^\w\s-]', '', data.topic)[:50].strip().replace(' ', '_')
    filename = f"synapse_{safe_name}.pdf"

    return Response(
        content=buffer.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


@app.post("/download-docx")
async def download_docx(data: DocxRequest):
    """Generate and return a Word (.docx) file from the report text."""
    if not data.report:
        return JSONResponse({"error": "No report provided"}, status_code=400)

    docx_bytes = markdown_to_docx(data.report, data.topic)
    buffer = io.BytesIO(docx_bytes)
    buffer.seek(0)

    safe_name = re.sub(r'[^\w\s-]', '', data.topic)[:50].strip().replace(' ', '_')
    filename = f"synapse_{safe_name}.docx"

    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


@app.get("/history")
async def get_history():
    """Return list of past research (summary only)."""
    history = load_history()
    return [
        {"id": h["id"], "topic": h["topic"], "timestamp": h["timestamp"]}
        for h in history
    ]


@app.get("/history/{entry_id}")
async def get_history_entry(entry_id: str):
    """Return full details of a specific research entry."""
    history = load_history()
    for h in history:
        if h["id"] == entry_id:
            return h
    return JSONResponse({"error": "Not found"}, status_code=404)


@app.delete("/history/{entry_id}")
async def delete_history_entry(entry_id: str):
    """Delete a specific research entry."""
    history = load_history()
    history = [h for h in history if h["id"] != entry_id]
    save_history(history)
    return {"ok": True}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("app:app", host="127.0.0.1", port=5000, reload=True)
